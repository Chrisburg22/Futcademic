#!/usr/bin/env python3
"""Convert the 4 architecture Markdown files to .docx using python-docx.

Handles:
- Headings (#, ##, ###, ####)
- Tables (GFM pipe tables)
- Fenced code blocks (``` ... ```)
  * mermaid blocks are preserved as monospace text with a header note
- Bold/italic inline (**text**, *text*, `code`)
- Bullet / numbered lists
- Blank lines → paragraph breaks
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent

SOURCES = [
    "P03-SWA_SecD15_Team6.md",
    "SWAReviewCheckList_SecD15_Team6.md",
    "P03-SDD_SecD15_Team6_Functionality.md",
    "P03-CodeReview_SecD15_Team6_Individual.md",
    "P05-PIF_SecD15_Team_6_Futcamedic.md",
    "P-CM_SecD15_Team_6_ConfigurationManagement.md",
]


def set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_inline(paragraph, text: str):
    """Parse **bold**, *italic*, `code` inline markdown inside a paragraph."""
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for m in token_re.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            run.font.name = "Calibri"
        tok = m.group(0)
        if tok.startswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif tok.startswith("*"):
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = "Calibri"


def parse_table_rows(lines: list[str]) -> list[list[str]]:
    """Parse GFM pipe table lines. Lines include header, separator, data."""
    rows = []
    for line in lines:
        if re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line):
            continue  # separator
        # split on | but ignore leading/trailing empties from | at edges
        parts = [p.strip() for p in line.strip().strip("|").split("|")]
        rows.append(parts)
    return rows


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        tr = table.rows[ri]
        for ci in range(ncols):
            cell = tr.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            val = row[ci] if ci < len(row) else ""
            add_inline(p, val)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_bg(cell, "D9E2F3")
    doc.add_paragraph()


def add_code_block(doc: Document, lang: str, code_lines: list[str]):
    if lang.lower() == "mermaid":
        note = doc.add_paragraph()
        r = note.add_run("[Diagrama Mermaid — render en Markdown / VSCode / GitHub]")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for i, line in enumerate(code_lines):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)


def md_to_docx(md_path: Path, out_path: Path):
    doc = Document()
    # base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # fenced code block
        m = re.match(r"^```(\w*)\s*$", line)
        if m:
            lang = m.group(1) or ""
            i += 1
            buf = []
            while i < n and not re.match(r"^```\s*$", lines[i]):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            add_code_block(doc, lang, buf)
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            h = doc.add_heading(level=min(level, 4))
            add_inline(h, heading_text)
            i += 1
            continue

        # horizontal rule
        if re.match(r"^---+\s*$", line):
            doc.add_paragraph("_" * 60)
            i += 1
            continue

        # pipe table: detect a line starting with | and next line being separator
        if line.lstrip().startswith("|") and i + 1 < n and re.match(
            r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[i + 1]
        ):
            tbl_lines = []
            while i < n and lines[i].lstrip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            rows = parse_table_rows(tbl_lines)
            add_table(doc, rows)
            continue

        # image: ![alt](path)
        m = re.match(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if m:
            img_path = (HERE / m.group(2)).resolve()
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(img_path), width=Cm(13))
            else:
                p = doc.add_paragraph()
                r = p.add_run(f"[Imagen no encontrada: {m.group(2)}]")
                r.italic = True
            i += 1
            continue

        # bullet list
        m = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.6 + len(m.group(1)) // 2 * 0.3)
            add_inline(p, m.group(2))
            i += 1
            continue

        # numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(2))
            i += 1
            continue

        # blank line
        if not line.strip():
            i += 1
            continue

        # paragraph
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"OK  {out_path.name}  ({out_path.stat().st_size // 1024} KB)")


def main():
    for md in SOURCES:
        src = HERE / md
        if not src.exists():
            print(f"MISS {md}")
            continue
        out = src.with_suffix(".docx")
        md_to_docx(src, out)


if __name__ == "__main__":
    main()
